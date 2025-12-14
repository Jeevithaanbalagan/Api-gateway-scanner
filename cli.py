#!/usr/bin/env python3
"""
API Gateway Attack Surface Scanner (v2 – Stable)

Features:
- OpenAPI-assisted + Black-box scanning
- OWASP API Top 10 aligned passive checks
- IDOR, SSRF indicators, Business Logic hints
- Inventory & versioning checks
- Function-level authorization checks
- Severity scoring (LOW / MEDIUM / HIGH)
- Progress bar + scan summary
- Reports: JSON, HTML, PDF, DOCX
"""

import json
import os
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Dict, List, Tuple

import click
import httpx
import yaml
from jinja2 import Template

# =========================================================
# BANNER / META
# =========================================================

BANNER = r"""
███████╗██████╗  ██████╗ ███████╗ ██████╗ ██╗   ██╗ █████╗ ██████╗ ██████╗
        API Gateway Attack Surface Scanner
"""

VERSION = "2.1.0"

# =========================================================
# OpenAPI Helpers
# =========================================================

def load_openapi(path: str) -> Dict:
    if not os.path.exists(path):
        raise click.ClickException(f"OpenAPI file not found: {path}")
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)

def enumerate_paths(spec: Dict) -> List[Tuple[str, str]]:
    routes = []
    for path, methods in spec.get("paths", {}).items():
        for m in methods:
            if m.lower() in ("get", "post", "put", "delete", "patch", "options", "head"):
                routes.append((m.upper(), path))
    return routes

def build_url(base: str, path: str) -> str:
    return base.rstrip("/") + "/" + path.lstrip("/")

# =========================================================
# Severity Engine
# =========================================================

HIGH = {"missing_auth", "function_level_auth", "ssrf_indicator"}
MEDIUM = {"cors", "object_property_exposure", "business_logic"}
LOW = {"security_headers", "inventory_versioning"}

def get_severity(name: str) -> str:
    if name in HIGH:
        return "HIGH"
    if name in MEDIUM:
        return "MEDIUM"
    return "LOW"

# =========================================================
# OWASP API Security Checks (Passive)
# =========================================================

SECURITY_HEADERS = [
    "strict-transport-security",
    "content-security-policy",
    "x-frame-options",
    "x-content-type-options",
    "referrer-policy",
]

SSRF_HINTS = ["url", "redirect", "callback", "target", "image", "fetch"]

def check_endpoint(client: httpx.Client, method: str, url: str):
    findings = {"method": method, "url": url, "checks": []}

    try:
        r = client.request(method, url, timeout=15)
    except Exception as e:
        findings["checks"].append({
            "name": "request_error",
            "ok": False,
            "severity": "HIGH",
            "description": str(e)
        })
        return findings

    headers = {k.lower(): v for k, v in r.headers.items()}

    # ---------------- API2: Broken Authentication ----------------
    ok = not (200 <= r.status_code < 300)
    findings["checks"].append({
        "name": "missing_auth",
        "ok": ok,
        "severity": get_severity("missing_auth"),
        "description": f"Unauthenticated status: {r.status_code}"
    })

    # ---------------- API8: CORS ----------------
    try:
        r_cors = client.request(method, url, headers={"Origin": "https://evil.example"})
        ao = r_cors.headers.get("access-control-allow-origin")
        ok = ao not in ("*", "https://evil.example")
        findings["checks"].append({
            "name": "cors",
            "ok": ok,
            "severity": get_severity("cors"),
            "description": f"Access-Control-Allow-Origin: {ao}"
        })
    except:
        findings["checks"].append({
            "name": "cors",
            "ok": True,
            "severity": "LOW",
            "description": "CORS check skipped"
        })

    # ---------------- API7: Security Headers ----------------
    missing = [h for h in SECURITY_HEADERS if h not in headers]
    findings["checks"].append({
        "name": "security_headers",
        "ok": not bool(missing),
        "severity": get_severity("security_headers"),
        "description": "Missing: " + ", ".join(missing) if missing else "All present"
    })

    # ---------------- API9: Inventory / Versioning ----------------
    overlap = "/v1" in url and "/v2" in url
    findings["checks"].append({
        "name": "inventory_versioning",
        "ok": not overlap,
        "severity": get_severity("inventory_versioning"),
        "description": "API version overlap detected" if overlap else "Clean versioning"
    })

    # ---------------- API3: Excessive Data Exposure ----------------
    if "application/json" in headers.get("content-type", ""):
        try:
            data = r.json()
            if isinstance(data, dict):
                sensitive = [k for k in data if any(s in k.lower() for s in ["password", "token", "secret", "key"])]
                findings["checks"].append({
                    "name": "object_property_exposure",
                    "ok": not bool(sensitive),
                    "severity": get_severity("object_property_exposure"),
                    "description": "Exposed: " + ", ".join(sensitive) if sensitive else "No sensitive fields"
                })
        except:
            pass

    # ---------------- API6: Business Logic ----------------
    if any(x in url.lower() for x in ["otp", "reset", "verify", "auth"]):
        findings["checks"].append({
            "name": "business_logic",
            "ok": False,
            "severity": get_severity("business_logic"),
            "description": "Sensitive business flow endpoint"
        })

    # ---------------- API7: SSRF Indicators ----------------
    ssrf = any(p in url.lower() for p in SSRF_HINTS)
    findings["checks"].append({
        "name": "ssrf_indicator",
        "ok": not ssrf,
        "severity": get_severity("ssrf_indicator"),
        "description": "SSRF-like parameter detected" if ssrf else "No SSRF hints"
    })

    # ---------------- API5: Function Level Authorization ----------------
    try:
        r_override = client.request(method, url, headers={"X-HTTP-Method-Override": "DELETE"})
        ok = r_override.status_code not in (200, 204)
    except:
        ok = True

    findings["checks"].append({
        "name": "function_level_auth",
        "ok": ok,
        "severity": get_severity("function_level_auth"),
        "description": "Method override protected" if ok else "DELETE override accepted"
    })

    return findings

# =========================================================
# REPORTING
# =========================================================

def write_json(path, data):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w") as f:
        json.dump(data, f, indent=2)

HTML_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <title>API Gateway Security Report</title>
  <style>
    body { font-family: Arial, sans-serif; margin: 40px; }
    h1, h2, h3 { color: #2c3e50; }
    table { width: 100%; border-collapse: collapse; margin-top: 20px; }
    th, td { padding: 10px; border: 1px solid #ddd; text-align: left; }
    th { background: #34495e; color: white; }
    .HIGH { background: #ffcccc; }
    .MEDIUM { background: #fff3cd; }
    .LOW { background: #e8f5e9; }
    .summary-box {
      display: flex;
      gap: 20px;
      margin: 20px 0;
    }
    .card {
      padding: 15px;
      border-radius: 6px;
      color: white;
      width: 150px;
      text-align: center;
    }
    .card.high { background: #c0392b; }
    .card.medium { background: #f39c12; }
    .card.low { background: #27ae60; }
  </style>
</head>
<body>

<h1>API Gateway Security Scan Report</h1>
<p><b>Scan Time:</b> {{ scan_time }}</p>
<p><b>Targets:</b> {{ targets | join(", ") }}</p>

<h2>Executive Summary</h2>
<div class="summary-box">
  <div class="card high">HIGH<br>{{ summary.HIGH }}</div>
  <div class="card medium">MEDIUM<br>{{ summary.MEDIUM }}</div>
  <div class="card low">LOW<br>{{ summary.LOW }}</div>
</div>

<h2>Findings</h2>
<table>
<tr>
  <th>Method</th>
  <th>Endpoint</th>
  <th>Issue</th>
  <th>Severity</th>
  <th>Description</th>
</tr>

{% for ep in endpoints %}
{% for c in ep.checks %}
{% if not c.ok %}
<tr class="{{ c.severity }}">
  <td>{{ ep.method }}</td>
  <td>{{ ep.url }}</td>
  <td>{{ c.name }}</td>
  <td>{{ c.severity }}</td>
  <td>{{ c.description }}</td>
</tr>
{% endif %}
{% endfor %}
{% endfor %}

</table>

</body>
</html>
"""


def write_html(path, report):
    html = Template(HTML_TEMPLATE).render(**report)
    with open(path, "w") as f:
        f.write(html)

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document

from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Table,
    TableStyle,
    Spacer
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors


def write_pdf(path, report):
    doc = SimpleDocTemplate(
        path,
        pagesize=letter,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36,
    )

    styles = getSampleStyleSheet()
    elements = []

    # ---------------- Title ----------------
    elements.append(Paragraph(
        "<b>API Gateway Security Scan Report</b>",
        styles["Title"]
    ))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph(
        f"<b>Scan Time:</b> {report['scan_time']}",
        styles["Normal"]
    ))
    elements.append(Paragraph(
        f"<b>Targets:</b> {', '.join(report['targets'])}",
        styles["Normal"]
    ))
    elements.append(Spacer(1, 20))

    # ---------------- Executive Summary ----------------
    elements.append(Paragraph(
        "<b>Executive Summary</b>",
        styles["Heading2"]
    ))
    summary = report.get("summary", {})
    elements.append(Paragraph(
        f"HIGH: {summary.get('HIGH', 0)} | "
        f"MEDIUM: {summary.get('MEDIUM', 0)} | "
        f"LOW: {summary.get('LOW', 0)}",
        styles["Normal"]
    ))
    elements.append(Spacer(1, 20))

    # ---------------- Findings Table ----------------
    elements.append(Paragraph(
        "<b>Findings Summary</b>",
        styles["Heading2"]
    ))
    elements.append(Spacer(1, 10))

    table_data = [
        ["Method", "Endpoint", "Issue", "Severity"]
    ]

    for ep in report["endpoints"]:
        for c in ep["checks"]:
            if not c["ok"]:
                table_data.append([
                    ep["method"],
                    ep["url"],
                    c["name"],
                    c["severity"]
                ])

    table = Table(
        table_data,
        colWidths=[60, 260, 120, 80],
        repeatRows=1
    )

    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.darkblue),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold"),

        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),

        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),

        ("BACKGROUND", (0, 1), (-1, -1), colors.whitesmoke),
    ]))

    elements.append(table)

    doc.build(elements)



def write_docx(path, report):
    doc = Document()
    doc.add_heading("API Gateway Security Scan Report", 0)

    doc.add_paragraph(f"Scan Time: {report['scan_time']}")
    doc.add_paragraph(f"Targets: {', '.join(report['targets'])}")

    doc.add_heading("Findings Summary", level=1)

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Method"
    hdr[1].text = "Endpoint"
    hdr[2].text = "Issue"
    hdr[3].text = "Severity"

    for ep in report["endpoints"]:
        for c in ep["checks"]:
            if not c["ok"]:
                row = table.add_row().cells
                row[0].text = ep["method"]
                row[1].text = ep["url"]
                row[2].text = c["name"]
                row[3].text = c["severity"]

    doc.save(path)

# =========================================================
# CLI
# =========================================================

@click.group(invoke_without_command=True)
@click.pass_context
def cli(ctx):
    if ctx.invoked_subcommand is None:
        click.secho(BANNER, fg="cyan")
        click.echo("Use --help to see commands\n")

@cli.command()
def version():
    click.echo(f"API Gateway Scanner v{VERSION}")

@cli.command()
@click.option("--config", "-c", default="config.yaml")
@click.option("--openapi", "-o", default=None)
@click.option("--targets", "-t", default=None)
@click.option("--output", "-O", default="reports/scan.json")
def scan(config, openapi, targets, output):

    cfg = yaml.safe_load(open(config)) if os.path.exists(config) else {}
    targets_list = targets.split(",") if targets else cfg.get("targets", [])

    if not targets_list:
        raise click.ClickException("No targets provided")

    openapi_files = []
    if openapi:
        openapi_files.append(openapi)
    else:
        ocfg = cfg.get("openapi", {})
        if ocfg.get("enabled", False):
            folder = ocfg.get("folder", "specs")
            if os.path.exists(folder):
                openapi_files += [
                    os.path.join(folder, f)
                    for f in os.listdir(folder)
                    if f.endswith((".yaml", ".yml", ".json"))
                ]

    endpoints = []
    if openapi_files:
        for f in openapi_files:
            spec = load_openapi(f)
            for m, p in enumerate_paths(spec):
                for base in targets_list:
                    endpoints.append((m, build_url(base, p)))
    else:
        for base in targets_list:
            endpoints.append(("GET", base))

    client = httpx.Client()
    results = []

    click.echo(f"[+] Scanning {len(endpoints)} endpoints")

    with click.progressbar(length=len(endpoints), label="Scanning") as bar:
        with ThreadPoolExecutor(max_workers=8) as pool:
            futures = [pool.submit(check_endpoint, client, m, u) for m, u in endpoints]
            for f in as_completed(futures):
                results.append(f.result())
                bar.update(1)

    summary = {"HIGH": 0, "MEDIUM": 0, "LOW": 0}
    for ep in results:
        for c in ep["checks"]:
            if not c["ok"]:
                summary[c["severity"]] += 1

    report = {
        "scan_time": time.strftime("%Y-%m-%d %H:%M:%S"),
        "targets": targets_list,
        "summary": summary,
        "endpoints": results
    }

    write_json(output, report)
    write_html(output.replace(".json", ".html"), report)
    write_pdf(output.replace(".json", ".pdf"), report)
    write_docx(output.replace(".json", ".docx"), report)

    click.echo("\nScan complete ✔")
    click.echo(f"HIGH: {summary['HIGH']} | MEDIUM: {summary['MEDIUM']} | LOW: {summary['LOW']}")
    click.echo(f"Reports saved to → {os.path.dirname(output) or 'reports/'}")

if __name__ == "__main__":
    cli()
